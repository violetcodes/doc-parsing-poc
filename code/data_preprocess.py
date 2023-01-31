import os
import docx
import string

def test_docx_module_on_folder(folder_path): #goes into utilities
    """check if docx module is able to open files in folder_path"""
    import docx
    from tqdm import tqdm
    error_files = []
    training_files_path = get_files_path(folder_path)
    i, j = 0, 0
    for doc_path in tqdm(training_files_path):
        try:
            document = docx.Document(doc_path)
            for paragraphs in document.paragraphs:
                paragraphs.text
        except:
            error_files(doc_path)
            j += 1
        i += 1

    print(f'Successfully parsed {i-j} documents out of {i} documents')
    
    return error_files

def get_files_path(folder_path): # subfolder is either 
    '''get function for getting list of file paths'''
    for i in os.listdir(folder_path):
        path_to_return = os.path.join(folder_path, i)
        if os.path.isfile(path_to_return):
            yield path_to_return

def get_full_text(file_path):
    "open doc files with docx"
    document = docx.Document(file_path)
    return '\n\n'.join([para.text for para in document.paragraphs])

# clean_text
# write rules
### cleaning rules section start ###

def remove_repeated_punctuations(text):
    for i in string.punctuation:
        while i*2 in text:
            text = text.replace(i*2, i)
    return text

def remove_multiple_next_lines(text):
    while '\n\n\n' in text:
        text = text.replace('\n\n\n', '\n\n')    
    return text

def remove_noisy_puctuations(text):
    punct_to_keep = ',.:)(-/' 
    punct_to_remove = [i for i in string.punctuation if i not in punct_to_keep]
    all_chars = list(set(text))
    for i in all_chars:
        if ord(i) > 126 or i in punct_to_remove:
            text = text.replace(i, '')
    return text

def remove_multiple_spaces(text):
    while '  ' in text:
        text = text.replace('  ', ' ')
    return text

def remove_tab_with_four_spaces(text):
    while '\t' in text:
        text = text.replace('\t', '    ')    
    return text

def merge_shorter_paragraphs(text):
    paragraphs = text.split('\n\n')
    final_paragraphs = []
    buffer = ''
    for para in paragraphs:
        
        para = para.replace('\n', ' ')
        if len(para) < 40:
            buffer += ' ' + para
        
        else:
            buffer += ' ' + para
            final_paragraphs.append(buffer.strip())
            buffer = ''
            
    return '\n\n'.join(final_paragraphs)

### cleaning rules section ends ###

def clean_text(text):    
    text = remove_repeated_punctuations(text)
    text = remove_multiple_next_lines(text)
    text = remove_noisy_puctuations(text)
    text = remove_multiple_spaces(text)
    text = remove_tab_with_four_spaces(text)
    
    text = merge_shorter_paragraphs(text)    
    return text.strip()



### Noramlize labels
def is_valid_date(day, month, year):
    import datetime
    try:
        datetime.datetime(year=year,month=month,day=day)
        return True
    except ValueError:
        return False

def normalize_date(date):
    if str(date) == 'nan':
        return None
    day, month, year = list(map(int, date.split('.')))
    if is_valid_date(day, month, year):
        return [day, month, year]    
    elif is_valid_date(day-1, month, year):
        return [day-1, month, year]
    elif is_valid_date(day-3, month, year):
        return [day-1, month, year]
    else: return None


def super_clean(text):
    # remove all_punctuations, and lower_case, double spaces
    paragraphs = [i for i in text.split('\n')]
    new_paragraphs = []
    for paragraph in paragraphs:
        paragraph = ''.join([i for i in paragraph if i not in string.punctuation])
        while '  ' in paragraph:
            paragraph = paragraph.replace('  ', ' ')
        new_paragraphs.append(paragraph.lower())
    return '\n\n'.join(new_paragraphs)  